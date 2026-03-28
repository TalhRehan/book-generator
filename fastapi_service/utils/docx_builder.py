from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_docx(title: str, chapters: list[dict], output_path: str):
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()
    doc.add_page_break()

    # table of contents header
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for ch in chapters:
        toc_line = doc.add_paragraph(style="List Number")
        toc_line.add_run(f"Chapter {ch['chapter_number']}: {ch['title']}")

    doc.add_page_break()

    # chapters
    for ch in chapters:
        heading = doc.add_heading(f"Chapter {ch['chapter_number']}: {ch['title']}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for para_text in ch["content"].split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue

            # treat lines starting with ## as subheadings
            if para_text.startswith("##"):
                subheading = para_text.replace("##", "").strip()
                doc.add_heading(subheading, level=2)
            else:
                para = doc.add_paragraph(para_text)
                para.paragraph_format.space_after = Pt(8)
                for run in para.runs:
                    run.font.size = Pt(11)

        doc.add_page_break()

    doc.save(output_path)